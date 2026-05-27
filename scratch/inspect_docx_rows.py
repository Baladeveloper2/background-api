import docx

docx_path = r"d:\project\frontend\public\assets\CL-2026-05-0010 - Integra Software (1).docx"
doc = docx.Document(docx_path)

print(f"Total tables in document: {len(doc.tables)}")

for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx} - Total Rows: {len(table.rows)}")
    for r_idx, row in enumerate(table.rows):
        row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        # Print first few chars of each cell
        print(f"  Row {r_idx}: {row_text[:6]}")
